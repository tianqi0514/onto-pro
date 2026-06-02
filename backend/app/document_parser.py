from __future__ import annotations

from pathlib import Path
from typing import Dict, List


SUPPORTED_EXTENSIONS = {".md", ".txt", ".docx", ".xlsx", ".pdf"}


def parse_document(path: Path) -> Dict[str, object]:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        text = path.read_text(encoding="utf-8", errors="replace")
    elif suffix == ".docx":
        text = _parse_docx(path)
    elif suffix == ".xlsx":
        text = _parse_xlsx(path)
    elif suffix == ".pdf":
        text = _parse_pdf(path)
    else:
        text = ""

    return {
        "path": str(path),
        "name": path.name,
        "extension": suffix,
        "supported": suffix in SUPPORTED_EXTENSIONS,
        "text": text,
        "char_count": len(text),
        "snippets": _snippets(text),
    }


def _parse_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    table_lines: List[str] = []
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(values):
                table_lines.append(" | ".join(values))
    return "\n".join(paragraphs + table_lines)


def _parse_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(str(path), read_only=True, data_only=True)
    lines: List[str] = []
    for sheet in workbook.worksheets:
        lines.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(max_row=120, values_only=True):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                lines.append(" | ".join(values))
    return "\n".join(lines)


def _parse_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    lines: List[str] = []
    for index, page in enumerate(reader.pages[:20], start=1):
        text = page.extract_text() or ""
        if text.strip():
            lines.append(f"# Page {index}\n{text.strip()}")
    return "\n".join(lines)


def _snippets(text: str) -> List[str]:
    cleaned = " ".join(text.split())
    if not cleaned:
        return []
    return [cleaned[index:index + 220] for index in range(0, min(len(cleaned), 660), 220)]
